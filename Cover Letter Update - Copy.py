from docx import Document
import datetime
import tkinter as tk
from tkinter import simpledialog

PATH = r'Path\to\template'
FINPATH = r'Path\to\final'

def getText(filename, date, position, company, site):
  doc = Document(filename)
  fullText = []
  for para in doc.paragraphs:
    fullText.append(para.text)
  for i in range(0,len(fullText)):
    fullText[i]= fullText[i].replace('[DATE]', date)
    fullText[i]= fullText[i].replace('[POSITION]', position)
    fullText[i]= fullText[i].replace('[COMPANY]', company)
    fullText[i]= fullText[i].replace('[JOBSITE]', site)
    #print(fullText[i])
  return fullText

ROOT = tk.Tk()

ROOT.withdraw()
# the input dialog
position = simpledialog.askstring(title="Position", prompt="Position?:")
company = simpledialog.askstring(title="Company", prompt="Company?:")
site = "Linkedin"
						



date = datetime.datetime.now()

month = date.strftime("%B")
day = date.strftime("%d")
year = date.strftime("%Y")

date = (month +" "+ day.lstrip("0") +", " + year)
#print(date)


text = getText(PATH, date, position, company, site)
print(text)

document = Document(FINPATH)
for i in document.paragraphs:
  i.text=''
for i in range(0,len(text)):
  document.paragraphs[i].text = text[i]
document.save(FINPATH)